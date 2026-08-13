from __future__ import annotations

import json
from pathlib import Path

from docx import Document


ROOT = Path(__file__).parent
OUTPUT = ROOT / "calibration-fixtures"

DISTRACTOR_TOPICS = [
    "badge color", "meeting room", "training calendar", "contact channel", "cover sheet",
    "print layout", "approval stamp", "folder icon", "desk location", "notification tone",
    "template font", "support queue", "archive label", "reviewer initials",
]


def main() -> None:
    sources = json.loads((ROOT / "calibration_sources.json").read_text(encoding="utf-8"))
    OUTPUT.mkdir(parents=True, exist_ok=True)
    for filename, paragraphs in sources.items():
        document = Document()
        document.add_heading(paragraphs[0], level=1)
        for paragraph in paragraphs[1:]:
            document.add_paragraph(paragraph)
        document.save(OUTPUT / filename)
    for language in ("en", "ar"):
        for index, topic in enumerate(DISTRACTOR_TOPICS, start=1):
            document = Document()
            if language == "en":
                document.add_heading(f"Blue Lantern operational note {index}", level=1)
                document.add_paragraph(
                    f"This Blue Lantern note concerns the {topic}. It contains no records-lifecycle rule."
                )
            else:
                document.add_heading(f"ملاحظة تشغيلية للفانوس الأزرق {index}", level=1)
                document.add_paragraph(
                    f"تتناول ملاحظة الفانوس الأزرق رقم {index} موضوعًا تشغيليًا فقط ولا تحدد قاعدة لدورة حياة السجلات."
                )
            document.save(OUTPUT / f"cal-{language}-distractor-{index:02d}.docx")
    print(f"Built {len(sources) + 2 * len(DISTRACTOR_TOPICS)} local calibration fixtures in {OUTPUT}")


if __name__ == "__main__":
    main()
