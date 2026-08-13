from io import BytesIO

from crag.api import create_app
from crag.config import Settings
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from pypdf import PdfWriter


def docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_heading("Reference", level=1)
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def blank_pdf_bytes() -> bytes:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buffer = BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


def make_client(tmp_path) -> TestClient:
    settings = Settings(data_dir=tmp_path, runtime="deterministic")
    return TestClient(create_app(settings))


def test_grounded_query_returns_valid_citations(tmp_path) -> None:
    with make_client(tmp_path) as client:
        workspace = client.post("/api/workspaces", json={"name": "Egypt research"}).json()
        upload = client.post(
            f"/api/workspaces/{workspace['id']}/documents",
            files={
                "file": (
                    "reference.docx",
                    docx_bytes("The capital of Egypt is Cairo. Cairo is located along the Nile River."),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"ocr": "false"},
        )
        assert upload.status_code == 201
        assert upload.json()["status"] == "ready"

        queued = client.post(
            f"/api/workspaces/{workspace['id']}/queries",
            json={"question": "What is the capital of Egypt?"},
        )
        assert queued.status_code == 202
        run = client.get(f"/api/runs/{queued.json()['id']}").json()
        assert run["status"] == "completed"
        assert run["result"]["disposition"] == "answered"
        citation_ids = {citation["id"] for citation in run["result"]["citations"]}
        assert citation_ids
        assert all(
            set(claim["citation_ids"]) <= citation_ids for claim in run["result"]["claims"]
        )


def test_unanswerable_query_refuses_after_bounded_correction(tmp_path) -> None:
    with make_client(tmp_path) as client:
        workspace = client.post("/api/workspaces", json={"name": "Policy"}).json()
        client.post(
            f"/api/workspaces/{workspace['id']}/documents",
            files={"file": ("policy.docx", docx_bytes("The retention period is seven years."))},
            data={"ocr": "false"},
        )
        queued = client.post(
            f"/api/workspaces/{workspace['id']}/queries",
            json={"question": "What is the author's private home address?"},
        ).json()
        run = client.get(f"/api/runs/{queued['id']}").json()
        assert run["status"] == "refused"
        assert run["correction_count"] == 1
        assert run["result"]["claims"] == []


def test_blank_pdf_surfaces_optional_ocr_state(tmp_path) -> None:
    with make_client(tmp_path) as client:
        workspace = client.post("/api/workspaces", json={"name": "Scans"}).json()
        upload = client.post(
            f"/api/workspaces/{workspace['id']}/documents",
            files={"file": ("scan.pdf", blank_pdf_bytes(), "application/pdf")},
            data={"ocr": "false"},
        )
        assert upload.status_code == 201
        assert upload.json()["status"] == "needs_ocr"
        assert "OCR" in upload.json()["error"]

        retry = client.post(
            f"/api/workspaces/{workspace['id']}/documents",
            files={"file": ("scan.pdf", blank_pdf_bytes(), "application/pdf")},
            data={"ocr": "true"},
        )
        assert retry.status_code == 201
        assert retry.json()["id"] == upload.json()["id"]
        assert retry.json()["status"] == "failed"
        assert retry.json()["error"]


def test_invalid_file_is_rejected(tmp_path) -> None:
    with make_client(tmp_path) as client:
        workspace = client.post("/api/workspaces", json={"name": "Safety"}).json()
        response = client.post(
            f"/api/workspaces/{workspace['id']}/documents",
            files={"file": ("notes.txt", b"not a supported document", "text/plain")},
            data={"ocr": "false"},
        )
        assert response.status_code == 422
